import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging
# PDF parsing
import PyPDF2
# DOCX parsing
from docx import Document
# File type detection
# from config import SUPPORTED_FORMATS, MAX_FILE_SIZE_MB
from django.conf import settings

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    
    def __init__(self):
        self.supported_formats = settings.SUPPORTED_FORMATS
        self.max_file_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
    
    def parse_resume(self, file_path: str) -> Dict[str, any]:
        file_path = Path(file_path)
        
        # Validate file
        validation_result = self._validate_file(file_path)
        if not validation_result['valid']:
            return {
                'success': False,
                'error': validation_result['error'],
                'filename': file_path.name,
                'text': '',
                'metadata': {}
            }
        
        try:
            # Extract text based on file extension
            text = self._extract_text(file_path)
            
            # Clean and process text
            cleaned_text = self._clean_text(text)
            
            # Extract basic metadata
            metadata = self._extract_metadata(cleaned_text, file_path)
            
            return {
                'success': True,
                'filename': file_path.name,
                'text': cleaned_text,
                'metadata': metadata,
                'file_size': file_path.stat().st_size,
                'file_type': file_path.suffix.lower()
            }
            
        except Exception as e:
            logger.error(f"Error parsing {file_path.name}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'filename': file_path.name,
                'text': '',
                'metadata': {}
            }
    
    def _validate_file(self, file_path: Path) -> Dict[str, any]:
        
        
        if not file_path.exists():
            return {'valid': False, 'error': 'File does not exist'}
        
        if not file_path.is_file():
            return {'valid': False, 'error': 'Path is not a file'}
        
        if file_path.suffix.lower() not in self.supported_formats:
            return {'valid': False, 'error': f'Unsupported format. Supported: {self.supported_formats}'}
        
        if file_path.stat().st_size > self.max_file_size:
            return {'valid': False, 'error': f'File too large. Max size: {settings.MAX_FILE_SIZE_MB}MB'}
        
        return {'valid': True, 'error': None}
    
    def _extract_text(self, file_path: Path) -> str:
        
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif extension == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF parsing")
        
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                    
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        
        return text
    
    def _extract_from_docx(self, file_path: Path) -> str:
        
        if Document is None:
            raise ImportError("python-docx is required for DOCX parsing")
        
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
        
        return text
    
    def _extract_from_txt(self, file_path: Path) -> str:
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise Exception(f"Error reading TXT file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\/\@\+]', '', text)
        
        # Remove multiple consecutive punctuation
        text = re.sub(r'[\.]{2,}', '.', text)
        text = re.sub(r'[\-]{2,}', '-', text)
        
        return text.strip()
    
    def _extract_metadata(self, text: str, file_path: Path) -> Dict[str, any]:
        
        metadata = {
            'word_count': len(text.split()),
            'char_count': len(text),
            'has_email': bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)),
            'has_phone': bool(re.search(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)),
            'has_linkedin': 'linkedin' in text.lower(),
            'has_github': 'github' in text.lower(),
        }

        # Extract candidate name with improved logic
        extracted_name = self._extract_candidate_name(text)
        metadata['candidate_name'] = extracted_name

        # Keep potential_name for backward compatibility
        metadata['potential_name'] = extracted_name

        return metadata

    def _extract_candidate_name(self, text: str) -> str:
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        # Enhanced patterns for name extraction
        name_patterns = [
            # Standard name patterns
            r'^([A-Z][a-zA-Z]+\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)?)\s*$',
            r'^([A-Z][a-zA-Z]+\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)?)\s*[,\|]',
            r'^([A-Z][a-zA-Z]+\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)?)\s*\n',
            # Name with middle initial
            r'^([A-Z][a-zA-Z]+\s+[A-Z]\.?\s+[A-Z][a-zA-Z]+)\s*',
            # Name in various formats
            r'([A-Z][a-zA-Z]+\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)?)\s*(?:Resume|CV|Profile)?$'
        ]

        # Check first 10 lines for name
        for i, line in enumerate(lines[:10]):
            # Skip lines that are clearly not names
            skip_keywords = ['email', 'phone', 'address', 'linkedin', 'github', 'experience',
                           'education', 'skills', 'objective', 'summary', 'profile', 'contact']

            if any(keyword in line.lower() for keyword in skip_keywords):
                continue

            # Skip lines with numbers, symbols, or too many words
            if re.search(r'[\d@\(\)\-\+]', line) or len(line.split()) > 4:
                continue

            for pattern in name_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    name = match.group(1).strip()
                    name_parts = name.split()

                    # Validate name parts
                    if (2 <= len(name_parts) <= 4 and
                        all(part.replace('.', '').isalpha() for part in name_parts) and
                        all(len(part) >= 2 for part in name_parts if '.' not in part)):

                        # Proper case the name
                        formatted_name = ' '.join(part.title() for part in name_parts)
                        return formatted_name

        # Fallback: check first few words of the entire text
        words = text.split()
        for i in range(min(10, len(words) - 1)):
            if (words[i].isalpha() and words[i+1].isalpha() and
                words[i][0].isupper() and words[i+1][0].isupper() and
                len(words[i]) >= 2 and len(words[i+1]) >= 2):

                # Check if it's not a common non-name word
                non_names = ['Dear', 'To', 'From', 'Subject', 'Date', 'Re', 'Regarding']
                if words[i] not in non_names and words[i+1] not in non_names:
                    return f"{words[i].title()} {words[i+1].title()}"

        return "Name not found"


# def test_parser():
#     parser = ResumeParser()
    
#     # Test with a sample text file
#     test_file = Path("test_resume.txt")
#     test_content = """
#     John Doe
#     Software Developer
#     Email: john.doe@email.com
#     Phone: (555) 123-4567
    
#     Experience:
#     - 3 years Python development
#     - React and JavaScript
#     - SQL databases
    
#     Education:
#     Bachelor's in Computer Science
#     """
    
#     # Create test file
#     with open(test_file, 'w') as f:
#         f.write(test_content)
    
#     # Parse test file
#     result = parser.parse_resume(test_file)
#     print("Test Result:", result)
    
#     # Clean up
#     test_file.unlink()


# if __name__ == "__main__":
#     test_parser()
